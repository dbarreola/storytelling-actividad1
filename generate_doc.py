#!/usr/bin/env python3
"""Generate A1_E4.docx for Storytelling activity."""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

doc = Document()

# --- Page margins ---
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# --- Styles ---
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(12)

# --- PORTADA ---
for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Universidad del Valle de México')
run.font.size = Pt(14)
run.font.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Solucionar para Cambiar')
run.font.size = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('SGTI0103A-531XO04B2602 · Unidad 1')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x6E, 0x66, 0x5C)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Actividad 1 — Storytelling')
run.font.size = Pt(20)
run.font.bold = True
run.font.color.rgb = RGBColor(0xB9, 0x2C, 0x32)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('El poder de las narrativas de marca y las historias de los consumidores')
run.font.size = Pt(12)
run.font.italic = True

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Equipo 4')
run.font.size = Pt(13)
run.font.bold = True

team = [
    'Diego Armando Beltrán Arreola',
    'Enrique Varela Campuzano',
    'Jorge Ignacio Valenzuela Padrón',
    'Luis Daniel Rosales Lozano',
    'Naomi Giselle Silva Morán',
    'Nelson Andrés Gutiérrez Morán',
]
for name in team:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(name)
    run.font.size = Pt(11)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Docente: Paola Argüelles Gómez')
run.font.size = Pt(11)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Fecha de entrega: 1 de junio de 2026')
run.font.size = Pt(11)

# --- PAGE BREAK ---
doc.add_page_break()

# --- LINK A PRESENTACIÓN ---
h = doc.add_heading('Enlace a la presentación interactiva', level=1)
p = doc.add_paragraph()
run = p.add_run('La presentación interactiva completa se encuentra disponible en el siguiente enlace:')
p = doc.add_paragraph()
run = p.add_run('https://dbarreola.github.io/storytelling-actividad1/')
run.font.color.rgb = RGBColor(0x00, 0x71, 0xE3)
run.font.size = Pt(11)
p = doc.add_paragraph()
run = p.add_run('Instrucciones: Abra el enlace en cualquier navegador. Use las teclas → ↓ o Espacio para navegar entre diapositivas.')
run.font.size = Pt(10)
run.font.italic = True
run.font.color.rgb = RGBColor(0x6E, 0x66, 0x5C)

doc.add_paragraph()

# --- SECCIÓN 1: EXPLORACIÓN DE MARCAS ---
doc.add_heading('1. Exploración de marcas', level=1)

doc.add_heading('1.1 Apple — Marca favorita', level=2)
doc.add_paragraph(
    'Apple es una empresa de tecnología fundada en 1976, reconocida por su diseño, '
    'ecosistema integrado y narrativa premium. Su storytelling se construye sobre cinco pilares: '
    '"Think Different" (el usuario como visionario), empoderamiento ("Shot on iPhone" muestra lo que '
    'el usuario crea, no la cámara), minimalismo (el vacío comunica premium), ecosistema exclusivo '
    '(un club cerrado que refuerza la lealtad tribal), y experiencia sobre producto (tiendas como '
    '"plazas públicas", keynotes como teatro).'
)

doc.add_heading('Reseñas positivas', level=3)
doc.add_paragraph(
    'Apple obtiene un NPS de 72/100 en el sector tecnológico y primer lugar en satisfacción '
    'móvil según J.D. Power 2023 (866/1000). El 65% de usuarios de iPhone poseen al menos '
    'otro producto Apple, y la fiabilidad del hardware alcanza 4.5/5 en el Six Colors Report Card 2026. '
    'Los consumidores elogian consistentemente la calidad del hardware, especialmente el chip Apple Silicon.'
)

doc.add_heading('Reseñas negativas', level=3)
doc.add_paragraph(
    '"El soporte de Apple no ayuda a sus clientes una vez que compras el producto" '
    '(Steven, WA, mayo 2026, ConsumerAffairs). "Si alguna vez te cambias a Android, Apple te '
    'bloquea el acceso a todas tus cuentas" (Kathy G., OR, oct. 2024, Sitejabber). '
    'En ConsumerAffairs, el 85% de las reseñas son de 1 estrella. Las quejas principales: '
    'servicio post-venta deficiente, ecosistema cerrado percibido como trampa, y obsolescencia programada.'
)

doc.add_heading('1.2 Xiaomi — Marca desconocida', level=2)
doc.add_paragraph(
    'Xiaomi es una empresa china fundada en 2010 que se posiciona como democratizadora de la '
    'tecnología con precios accesibles. Su narrativa se construye sobre: "Innovación para todos" '
    '(especificaciones premium sin sobreprecio), "Human x Car x Home" (ecosistema conectado), '
    '"Honest pricing" (precio transparente y justo), comunidad co-creadora (foros, Mi Fan Festival, '
    'betas colaborativas), y su posición de challenger innovador ("Vinimos para quedarnos", '
    'Jia Wei, Forbes, 2026).'
)

doc.add_heading('Reseñas positivas', level=3)
doc.add_paragraph(
    '"Es una de las compañías tech más versátiles del mundo. Todo lo que he usado ha sido de '
    'buena calidad" (usuario verificado, SiteJabber). Los consumidores destacan la relación '
    'calidad-precio excepcional, la versatilidad del ecosistema, y la duración de batería. '
    'Xiaomi es la tercera marca de smartphones en Europa tras Apple y Samsung.'
)

doc.add_heading('Reseñas negativas', level=3)
doc.add_paragraph(
    '"La PEOR marca de smartphones... ¡NO DESPERDICIEN SU DINERO en Xiaomi!" (usuario verificado, '
    'SiteJabber). "Es una estafa. Tu producto Mi se vuelve obsoleto si necesita reparación. '
    'Falta de repuestos." En SiteJabber obtiene 3.1/5 con distribución polarizada: 44% dan 5 estrellas, '
    '41% dan 1 estrella. El servicio post-venta deficiente es la queja más consistente.'
)

# --- SECCIÓN 2: ANÁLISIS ---
doc.add_heading('2. Análisis de narrativas', level=1)

doc.add_heading('2.1 Apple — Reflexión', level=2)
doc.add_paragraph(
    'Las reseñas de Apple revelan una paradoja: en encuestas estructuradas (J.D. Power, NPS) la '
    'satisfacción es altísima, pero en plataformas abiertas (Trustpilot, ConsumerAffairs) predominan '
    'las quejas. La narrativa de hardware premium y diseño sí se alinea con la experiencia real. '
    'Donde la historia se quiebra es en el servicio post-venta y la promesa de "el cliente primero". '
    'Esto refuerza nuestra preferencia, pero con matices: admiramos la excelencia en producto y narrativa, '
    'reconociendo que su debilidad está donde la marca deja de contar historias y empieza a prestar servicio.'
)

doc.add_heading('2.2 Xiaomi — Percepción', level=2)
doc.add_paragraph(
    'Antes de investigar, percibíamos a Xiaomi como "la alternativa china barata". Después de leer '
    'reseñas y analizar su narrativa, nuestra percepción cambió: es una empresa ambiciosa en plena '
    'transición de marca económica a ecosistema premium (vehículo eléctrico SU7, colaboración con Leica). '
    'El problema: su historia no ha alcanzado a su producto. Los consumidores reconocen el valor pero '
    'no sienten conexión emocional. Historias que conectarían mejor: testimoniales reales de cómo su '
    'tecnología accesible cambió la vida de alguien concreto; narrativas de su servicio post-venta '
    'mejorando; mostrar la comunidad de Mi Fans como protagonistas, no como público.'
)

doc.add_heading('2.3 El papel de las historias en la reputación', level=2)
doc.add_paragraph(
    'Apple: Su storytelling es tan fuerte que sobrevive a reseñas negativas. La tribu ya está '
    'convencida; los detractores son percibidos como excepciones. La narrativa actúa como escudo '
    'reputacional. Xiaomi: Sin una narrativa emocional fuerte, cada reseña negativa pesa más. '
    'No hay "tribu" que defienda la marca. La ausencia de storytelling deja la reputación vulnerable al ruido.'
)

# --- SECCIÓN 3: HISTORIA ---
doc.add_heading('3. Historia original: La burbuja verde', level=1)

doc.add_heading('Conexión con la investigación', level=2)
doc.add_paragraph(
    'De nuestra investigación aprendimos que el storytelling de Apple no solo vende productos; '
    'construye identidad tribal. La "burbuja azul" de iMessage no es un accidente de diseño: es un '
    'marcador de pertenencia. Quien está dentro del ecosistema es parte del club; quien no, queda '
    'excluido con una burbuja verde. Xiaomi ofrece la misma funcionalidad a una fracción del precio, '
    'pero sin narrativa emocional que contrarreste el estigma. Esta historia aplica lo que aprendimos: '
    'las narrativas de marca tienen consecuencias reales en comunidades reales.'
)

doc.add_heading('Problemática', level=2)
doc.add_paragraph(
    'En las preparatorias y universidades de México, el teléfono celular dejó de ser una herramienta '
    'de comunicación para convertirse en un marcador social. Las ventas de iPhone en México crecieron '
    'un 80% en el primer trimestre de 2026, y con ese crecimiento llegó un fenómeno documentado '
    'internacionalmente como "green bubble shaming": la exclusión social de quienes no tienen iPhone.'
)
doc.add_paragraph(
    'La problemática es real: grupos de iMessage donde solo entran usuarios de Apple, AirDrop que solo '
    'funciona entre iPhones, FaceTime que deja fuera a quien tiene Android. No son limitaciones '
    'accidentales; son decisiones de diseño de una marca que construyó su narrativa sobre la exclusividad. '
    'Y en una comunidad de adolescentes donde pertenecer lo es todo, quedar fuera del chat equivale a '
    'quedar fuera del grupo.'
)

doc.add_heading('Trascendencia', level=2)
doc.add_paragraph(
    'El Children\'s National Hospital de Washington ha catalogado el "green bubble shaming" como una '
    'forma de cyberbullying que causa depresión, baja autoestima y aislamiento escolar. En México, '
    'donde el salario mínimo mensual equivale al costo de un iPhone, la presión social se multiplica. '
    'Las familias se endeudan para que sus hijos "no se queden fuera". Los afectados son los estudiantes '
    '(que sufren el estigma), las familias (que absorben la deuda), y la comunidad escolar (que normaliza '
    'la discriminación por marca). Deberían involucrarse las instituciones educativas, las propias marcas, '
    'y los padres de familia.'
)

doc.add_heading('Detalle', level=2)
doc.add_paragraph(
    'Valeria tiene 16 años y estudia en una preparatoria pública en Guadalajara. En agosto, su mamá '
    'le compró un Xiaomi Redmi Note 13 a 12 meses sin intereses. Es un buen teléfono: cámara de '
    '108 megapíxeles, batería que dura dos días, pantalla AMOLED. En especificaciones, supera a '
    'varios iPhones. Valeria estaba contenta.'
)
doc.add_paragraph(
    'El problema empezó en septiembre, cuando sus amigas crearon un grupo de iMessage para organizar '
    'la fiesta de quince años de Sofía. Valeria no recibió la invitación. No por olvido; es que '
    'iMessage no admite a quien no tiene iPhone. "Es que no te llega, Val, no es por ti", le dijeron. '
    'Pero el resultado fue el mismo: las fotos, los planes, los chistes internos sucedían en un espacio '
    'donde ella no existía.'
)
doc.add_paragraph(
    'En el salón, Valeria empezó a notar los detalles. Las compañeras haciendo AirDrop de apuntes '
    'entre sí. El "¿por qué tus mensajes llegan en verde?" dicho como broma pero sentido como '
    'señalamiento. La funda de su Xiaomi que empezó a guardar boca abajo sobre la mesa para que no '
    'se viera la marca.'
)

doc.add_heading('Emoción', level=2)
doc.add_paragraph(
    'Un jueves de noviembre, Valeria le pidió a su mamá un iPhone. No dijo "quiero un iPhone"; dijo '
    '"es que mi teléfono no sirve para los trabajos en equipo". Su mamá, que trabaja turno doble en '
    'una maquiladora, le preguntó qué tenía de malo el Xiaomi. Valeria no supo explicar que el '
    'teléfono funcionaba perfecto, pero que ella se sentía invisible.'
)
doc.add_paragraph(
    'Lo que rompió el ciclo no fue comprar un iPhone. Fue una maestra de comunicación que, al notar '
    'que tres alumnas quedaban fuera de los grupos de organización, exigió que toda comunicación del '
    'salón fuera por WhatsApp, donde el color de la burbuja no existe. Un gesto institucional, '
    'pequeño pero deliberado.'
)
doc.add_paragraph(
    'Valeria sigue con su Xiaomi. La fiesta de Sofía fue bonita; Valeria se enteró de los detalles '
    'el mismo día, no tres días después. La marca de su teléfono no cambió, pero dejó de ser motivo '
    'de vergüenza cuando el entorno dejó de premiar la exclusividad. El problema nunca fue el teléfono. '
    'Fue la historia que una marca nos convenció de creer: que pertenecer tiene precio, y que ese '
    'precio lo pone Apple.'
)

# --- SECCIÓN 4: RECURSOS VISUALES ---
doc.add_heading('4. Recursos visuales', level=1)
doc.add_paragraph(
    'Se diseñaron 2 recursos visuales utilizando herramientas de inteligencia artificial (ChatGPT) '
    'que complementan y enriquecen la historia presentada:'
)
doc.add_paragraph(
    '1. Infografía "La paradoja de la conexión digital": presenta tres estadísticas clave '
    '(5,000+ conexiones digitales promedio, 3.2 amigos cercanos en adultos relocalizados, '
    '60% de jóvenes profesionales tech reportan soledad frecuente) en formato editorial minimalista.'
)
doc.add_paragraph(
    '2. Diagrama "El ciclo de la exclusión y la inclusión": representación visual del arco narrativo '
    'en cuatro pasos (Mudanza/nueva ciudad → Rutina aislada → Gesto humano → Comunidad emergente) '
    'con códigos de color semánticos.'
)
doc.add_paragraph(
    'Ambos recursos se encuentran integrados en la presentación interactiva (sección 4).'
)

# --- SECCIÓN 5: REFLEXIÓN ---
doc.add_heading('5. Reflexión final', level=1)
doc.add_paragraph(
    'Esta investigación revela un patrón claro: las historias no son decoración para marcas; son '
    'infraestructura de reputación. Apple sobrevive reseñas devastadoras porque ha invertido décadas '
    'en construir una narrativa tan sólida que funciona como escudo. Xiaomi, con un producto '
    'competitivo, queda expuesta a cada crítica porque no ha construido ese capital narrativo.'
)
doc.add_paragraph(
    'Los comentarios de consumidores que realmente influyen comparten una estructura: son específicos '
    '(nombres, fechas, situaciones concretas), son emocionales sin ser histriónicos, y conectan la '
    'experiencia individual con algo universal. Una reseña que dice "el teléfono es malo" no mueve a '
    'nadie; una que dice "mi madre intentó videollamarme desde el hospital y la pantalla se congeló" '
    'cambia decisiones de compra.'
)
doc.add_paragraph(
    'Las historias de marca y las historias de consumidores operan en la misma frecuencia emocional: '
    'ambas traducen hechos en significado. La diferencia es que las marcas controlan su narrativa '
    'mientras que los consumidores la verifican o la desmienten con experiencia vivida. No elegimos '
    'productos; elegimos las historias que queremos contar sobre nosotros mismos. Apple lo entiende y '
    'lo monetiza. Xiaomi aún está aprendiendo que el mejor precio del mercado no basta si no hay una '
    'historia que le dé significado a la compra.'
)
doc.add_paragraph(
    'La historia de Valeria lo demuestra: Apple no necesitó decirle que su Xiaomi era inferior. La '
    'narrativa de exclusividad hizo el trabajo sola, filtrándose hasta el salón de clases de una prepa '
    'pública en Guadalajara. Esa es la verdadera medida del poder del storytelling: cuando la historia '
    'que cuenta una marca se vuelve la historia que una comunidad vive, para bien o para mal.'
)
doc.add_paragraph(
    'La lección para cualquier comunicador: los datos informan, pero las historias transforman. Y las '
    'historias más poderosas no son las que cuenta la marca, sino las que los consumidores eligen repetir.'
)

# --- REFERENCIAS ---
doc.add_heading('Referencias', level=1)

refs = [
    'ConsumerAffairs. (2026). Apple Store reviews. https://www.consumeraffairs.com/computers/apple-store.html',
    'Gottschall, J. (2013). The storytelling animal: How stories make us human. Mariner Books.',
    'Heath, C., y Heath, D. (2007). Made to stick: Why some ideas survive and others die. Random House.',
    'Jia, W. (2026, febrero). How Xiaomi is winning the European smartphone market. Forbes. https://www.forbes.com/sites/ewanspence/2026/02/16/how-xiaomi-is-winning-european-smartphone-market/',
    'Retently. (2024). Apple\'s customer loyalty and high NPS. Retently Blog. https://www.retently.com/blog/apple-nps/',
    'Salva, D. (2025, diciembre). Beyond the byte: A deep dive into Apple\'s masterful storytelling. Dan Salva.',
    'SiteJabber. (2026). Apple reviews. https://www.sitejabber.com/reviews/apple.com',
    'SiteJabber. (2026). Xiaomi reviews. https://www.sitejabber.com/reviews/mi.com',
    'Six Colors. (2026, febrero). Apple Report Card 2026. Six Colors. https://sixcolors.com/post/2026/02/apple-report-card-2026/',
    'TEDx Talks. (2017, 16 de marzo). The magical science of storytelling | David JP Phillips | TEDxStockholm [Video]. YouTube. https://www.youtube.com/watch?v=Nj-hdQMa3uA',
    'World Health Organization. (2023). Social isolation and loneliness. WHO Commission on Social Connection.',
    'Xiaomi. (2026). ET BrandEquity: Xiaomi bets on trust, local roots and an ecosystem future. Economic Times. https://brandequity.economictimes.indiatimes.com/',
]

for ref in refs:
    p = doc.add_paragraph(ref)
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.first_line_indent = Cm(-1.27)
    p.paragraph_format.space_after = Pt(6)

# --- Save ---
output = os.path.expanduser('~/Documents/A1_DBL/A1_E4.docx')
doc.save(output)
print(f'Document saved: {output}')
